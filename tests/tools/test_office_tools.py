from __future__ import annotations

import json
import subprocess
from pathlib import Path

import pytest

from tools import office_tools


def _parse(result: str) -> dict:
    return json.loads(result)


@pytest.fixture
def workspace(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.setenv("TERMINAL_CWD", str(tmp_path))
    monkeypatch.setenv("HERMES_HOME", str(tmp_path / "hermes-home"))
    return tmp_path


def _require_office_libs() -> None:
    for mod in ("docx", "openpyxl", "pptx"):
        pytest.importorskip(mod)


# --------------------------------------------------------------------------- argument validation

def test_office_word_requires_action() -> None:
    result = _parse(office_tools.office_word_tool({}))
    assert "error" in result
    assert "read_text" in result["supported_actions"]


def test_office_excel_requires_path_for_read_actions() -> None:
    result = _parse(office_tools.office_excel_tool({"action": "read_sheet"}))
    assert "error" in result


def test_office_powerpoint_requires_parameters() -> None:
    result = _parse(office_tools.office_powerpoint_tool({"action": "add_slide"}))
    assert "error" in result


def test_run_script_reports_missing_script(tmp_path: Path) -> None:
    missing = tmp_path / "does-not-exist.py"
    result = _parse(office_tools._run_script(missing, []))
    assert result["error"].startswith("Missing script:")


def test_action_normalization_variants() -> None:
    assert office_tools._normalize_action("Read Metadata") == "read_metadata"
    assert office_tools._normalize_action("read-metadata") == "read_metadata"
    assert office_tools._normalize_action("  READ  ") == "read"


def test_word_action_alias_convert_pdf_requires_path() -> None:
    result = _parse(office_tools.office_word_tool({"action": "convert-pdf"}))
    assert "convert requires path" in result["error"]


@pytest.mark.parametrize(
    ("handler", "action"),
    [
        (office_tools.office_word_tool, "generate_exec_report"),
        (office_tools.office_word_tool, "validate_report"),
        (office_tools.office_excel_tool, "generate_kpi_workbook"),
        (office_tools.office_excel_tool, "validate_workbook"),
        (office_tools.office_powerpoint_tool, "generate_review_deck"),
        (office_tools.office_powerpoint_tool, "validate_deck"),
        (office_tools.office_powerpoint_tool, "pack"),
    ],
)
def test_fake_generators_and_unpack_actions_are_gone(handler, action) -> None:
    result = _parse(handler({"action": action, "output_path": "x", "path": "x", "output_file": "x"}))
    assert result["error"].startswith("Unsupported")


def test_schemas_have_no_generator_actions() -> None:
    for schema in (
        office_tools.OFFICE_WORD_SCHEMA,
        office_tools.OFFICE_EXCEL_SCHEMA,
        office_tools.OFFICE_POWERPOINT_SCHEMA,
    ):
        actions = schema["parameters"]["properties"]["action"]["enum"]
        assert not any(a.startswith(("generate_", "validate_")) for a in actions)
        assert schema["parameters"]["additionalProperties"] is False


def test_missing_input_file_is_reported_before_running_script(workspace: Path) -> None:
    result = _parse(office_tools.office_word_tool({"action": "read_text", "path": "nope.docx"}))
    assert result["error"].startswith("File not found")
    assert str(workspace) in result["error"]


def test_write_to_sensitive_path_is_refused(workspace: Path) -> None:
    result = _parse(
        office_tools.office_word_tool(
            {"action": "from_markdown", "text": "# x", "output_path": "/etc/hermes-office-test.docx"}
        )
    )
    assert "Refusing to write" in result["error"]


def test_check_office_tools_false_when_lib_missing(monkeypatch: pytest.MonkeyPatch) -> None:
    import importlib.util

    real = importlib.util.find_spec

    def _fake(name, *a, **k):
        if name == "pptx":
            return None
        return real(name, *a, **k)

    monkeypatch.setattr(importlib.util, "find_spec", _fake)
    assert office_tools._check_office_tools() is False
    assert office_tools.missing_office_dependencies() == ["pptx"]


def test_script_timeout_is_reported(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    script = tmp_path / "slow.py"
    script.write_text("import time; time.sleep(5)\n", encoding="utf-8")
    monkeypatch.setattr(office_tools, "missing_office_dependencies", lambda: [])

    def _raise(*args, **kwargs):
        raise subprocess.TimeoutExpired(cmd=args[0], timeout=kwargs.get("timeout"))

    monkeypatch.setattr(office_tools.subprocess, "run", _raise)
    monkeypatch.setenv("HERMES_OFFICE_TIMEOUT", "7")
    result = _parse(office_tools._run_script(script, []))
    assert "timed out after 7s" in result["error"]


def test_script_timeout_env_is_clamped(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("HERMES_OFFICE_TIMEOUT", "1")
    assert office_tools._script_timeout() == 5
    monkeypatch.setenv("HERMES_OFFICE_TIMEOUT", "nonsense")
    assert office_tools._script_timeout() == office_tools._DEFAULT_TIMEOUT_SECONDS


def test_missing_dependency_error_has_install_hint(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    script = tmp_path / "s.py"
    script.write_text("print('hi')\n", encoding="utf-8")
    monkeypatch.setattr(office_tools, "missing_office_dependencies", lambda: ["docx"])
    result = _parse(office_tools._run_script(script, []))
    assert "docx" in result["error"]
    assert "[office]" in result["hint"]


def test_markdown_to_slide_sections_parses_headings_and_bullets() -> None:
    md = """# Intro
- one
- two

## Details
1. alpha
2. beta
"""
    sections = office_tools._markdown_to_slide_sections(md)
    assert sections[0] == ("Intro", ["one", "two"])
    assert sections[1] == ("Details", ["alpha", "beta"])
    assert office_tools._markdown_to_slide_sections("") == []


def test_temp_files_are_cleaned_up(workspace: Path) -> None:
    temps = office_tools._TempFiles()
    path = temps.write_text("hello", ".md")
    assert Path(path).exists()
    temps.cleanup()
    assert not Path(path).exists()
    assert not any(p.name.startswith("office_") for p in workspace.iterdir())


# --------------------------------------------------------------------------- real round trips

def test_word_roundtrip_find_replace_spans_runs(workspace: Path) -> None:
    _require_office_libs()
    from docx import Document

    created = _parse(
        office_tools.office_word_tool(
            {"action": "from_markdown", "text": "# Title\n\nHello Old Text here.", "output_path": "r.docx"}
        )
    )
    assert created["success"], created
    doc = Document(str(workspace / "r.docx"))
    para = doc.add_paragraph()
    para.add_run("Ol")
    bold = para.add_run("d Te")
    bold.bold = True
    para.add_run("xt tail")
    doc.save(str(workspace / "r.docx"))

    replaced = _parse(
        office_tools.office_word_tool(
            {"action": "find_replace", "path": "r.docx", "find_text": "Old Text", "replace_text": "New Text"}
        )
    )
    assert replaced["success"], replaced
    assert '"replaced": 2' in replaced["stdout"]

    read = _parse(office_tools.office_word_tool({"action": "read_text", "path": "r.docx"}))
    assert "Hello New Text here." in read["stdout"]
    assert "New Text tail" in read["stdout"]
    assert "Old Text" not in read["stdout"]
    assert not any(p.name.startswith("office_") for p in workspace.iterdir())


def test_excel_roundtrip_batch_csv_and_truncation(workspace: Path) -> None:
    _require_office_libs()

    created = _parse(office_tools.office_excel_tool({"action": "create", "output_path": "w.xlsx", "sheets": ["Data", "Summary"]}))
    assert created["success"], created

    batch = _parse(
        office_tools.office_excel_tool(
            {
                "action": "set_cells",
                "path": "w.xlsx",
                "sheet": "Data",
                "cells": {"A1": "Region", "B1": "Revenue", "A2": "North", "B2": 1200.5},
            }
        )
    )
    assert batch["success"], batch
    assert '"cells_set": 4' in batch["stdout"]

    csv_row = _parse(office_tools.office_excel_tool({"action": "append_row", "path": "w.xlsx", "sheet": "Data", "row_csv": 'South,"1,300"'}))
    assert csv_row["success"], csv_row
    typed_row = _parse(office_tools.office_excel_tool({"action": "append_row", "path": "w.xlsx", "sheet": "Data", "values": ["East", 42]}))
    assert typed_row["success"], typed_row

    styled = _parse(
        office_tools.office_excel_tool(
            {"action": "format_cells", "path": "w.xlsx", "sheet": "Data", "range": "A1:B1", "style": {"bold": True, "fill": "3F59FF", "autofit": True}}
        )
    )
    assert styled["success"], styled

    import openpyxl

    wb = openpyxl.load_workbook(str(workspace / "w.xlsx"))
    ws = wb["Data"]
    assert ws["A3"].value == "South" and ws["B3"].value == "1,300"
    assert ws["A4"].value == "East" and ws["B4"].value == 42
    assert ws["A1"].font.bold is True
    assert ws["A1"].fill.start_color.rgb.endswith("3F59FF")
    wb.close()

    block = _parse(office_tools.office_excel_tool({"action": "read_sheet", "path": "w.xlsx", "sheet": "Data", "range": "A1:B2"}))
    assert block["stdout"].splitlines()[0] == "Region\tRevenue"

    truncated = _parse(office_tools.office_excel_tool({"action": "read_sheet", "path": "w.xlsx", "sheet": "Data", "max_rows": 1}))
    assert "[truncated]" in truncated["stdout"]


def test_powerpoint_roundtrip_create_edit_read(workspace: Path) -> None:
    _require_office_libs()

    created = _parse(
        office_tools.office_powerpoint_tool(
            {
                "action": "create",
                "output_path": "d.pptx",
                "title": "Deck",
                "subtitle": "Sub",
                "slides": [{"title": "One", "bullets": ["a", {"text": "b", "level": 1}], "notes": "n1"}, "Two"],
            }
        )
    )
    assert created["success"], created

    added = _parse(office_tools.office_powerpoint_tool({"action": "add_slide", "path": "d.pptx", "title": "Inserted", "bullets": ["x Old Text y"], "index": 2}))
    assert added["success"], added
    assert '"added_at": 2' in added["stdout"]

    replaced = _parse(office_tools.office_powerpoint_tool({"action": "find_replace", "path": "d.pptx", "find_text": "Old Text", "replace_text": "New"}))
    assert '"replaced": 1' in replaced["stdout"]

    read = _parse(office_tools.office_powerpoint_tool({"action": "read_text", "path": "d.pptx", "format": "json"}))
    detail = json.loads(read["stdout"])["detail"]
    assert [s["title"] for s in detail] == ["Deck", "Inserted", "One", "Two"]
    assert detail[1]["text"][-1] == "x New y"
    assert detail[2]["notes"] == "n1"

    deleted = _parse(office_tools.office_powerpoint_tool({"action": "delete_slide", "path": "d.pptx", "index": 2}))
    assert '"slides": 3' in deleted["stdout"]

    from_md = _parse(
        office_tools.office_powerpoint_tool(
            {"action": "create", "output_path": "d2.pptx", "title": "MD", "text": "# A\n- 1\n# B\nline", "template": "d.pptx"}
        )
    )
    assert from_md["success"], from_md
    read2 = _parse(office_tools.office_powerpoint_tool({"action": "read_text", "path": "d2.pptx", "format": "json"}))
    assert [s["title"] for s in json.loads(read2["stdout"])["detail"]] == ["MD", "A", "B"]

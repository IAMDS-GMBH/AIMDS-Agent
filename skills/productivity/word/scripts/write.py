#!/usr/bin/env python3
"""Create and edit Word (.docx) files.

Usage:
    # Create from a Markdown or plain-text file
    python write.py --from-md input.md output.docx

    # Find & replace text in an existing document
    python write.py --replace "Old Text" "New Text" document.docx

    # Append a paragraph to an existing document
    python write.py --append "New paragraph text." document.docx

    # Merge multiple documents into one
    python write.py --merge doc1.docx doc2.docx doc3.docx --out merged.docx
"""

import copy
import sys


def from_markdown(src, dst):
    """Convert a Markdown / plain-text file to .docx via pandoc (preferred)
    or a simple paragraph-by-paragraph fallback."""
    import subprocess
    from pathlib import Path

    # Try pandoc first — best quality
    try:
        result = subprocess.run(["pandoc", src, "-o", dst], capture_output=True, text=True)
        if result.returncode == 0:
            print(f"Created {dst} via pandoc.")
            return
    except FileNotFoundError:
        pass

    # Fallback: plain paragraph write
    from docx import Document

    doc = Document()
    for line in Path(src).read_text(encoding="utf-8").splitlines():
        line = line.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line:
            doc.add_paragraph(line)
    doc.save(dst)
    print(f"Created {dst} (pandoc not found; used plain fallback).")


def _replace_in_paragraph(para, old: str, new: str) -> int:
    """Replace ``old`` with ``new`` inside one paragraph, spanning run boundaries.

    Word splits text into runs at every formatting/spell-check/edit boundary,
    so a naive per-run replace silently misses matches like ``Ol|d Text``.
    We locate matches on the joined paragraph text, then rewrite only the runs
    a match touches: the first run keeps its formatting and receives the
    replacement, intermediate runs are emptied, the last run keeps its tail.
    Runs outside any match are left untouched. Returns the number of matches.
    """
    runs = list(para.runs)
    if not runs or not old:
        return 0
    texts = [r.text for r in runs]
    full = "".join(texts)
    if old not in full:
        return 0

    positions = []
    i = full.find(old)
    while i != -1:
        positions.append(i)
        i = full.find(old, i + len(old))

    bounds = []
    offset = 0
    for t in texts:
        bounds.append((offset, offset + len(t)))
        offset += len(t)

    def _run_at(pos: int, *, end: bool) -> int:
        for idx, (start, stop) in enumerate(bounds):
            if start == stop:
                continue
            if end:
                if start < pos <= stop:
                    return idx
            elif start <= pos < stop:
                return idx
        return len(bounds) - 1

    # Process from the last match backwards so earlier offsets stay valid.
    for pos in reversed(positions):
        match_end = pos + len(old)
        first = _run_at(pos, end=False)
        last = _run_at(match_end, end=True)
        a = pos - bounds[first][0]
        b = match_end - bounds[last][0]
        if first == last:
            texts[first] = texts[first][:a] + new + texts[first][b:]
        else:
            texts[first] = texts[first][:a] + new
            for mid in range(first + 1, last):
                texts[mid] = ""
            texts[last] = texts[last][b:]

    for run, text in zip(runs, texts):
        if run.text != text:
            run.text = text
    return len(positions)


def _iter_paragraphs(doc):
    """Yield every paragraph in body, tables (nested), headers and footers."""

    def _from_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from _from_tables(cell.tables)

    yield from doc.paragraphs
    yield from _from_tables(doc.tables)
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            yield from part.paragraphs
            yield from _from_tables(part.tables)


def find_replace(old, new, path):
    import json

    from docx import Document

    doc = Document(path)
    count = 0
    paragraphs_changed = 0
    for para in _iter_paragraphs(doc):
        n = _replace_in_paragraph(para, old, new)
        if n:
            count += n
            paragraphs_changed += 1

    if count:
        doc.save(path)
    print(f"Replaced {count} occurrence(s) of {old!r} -> {new!r} in {path}")
    print(json.dumps({"replaced": count, "paragraphs_changed": paragraphs_changed, "file": str(path)}))


def append_paragraph(text, path):
    from docx import Document

    doc = Document(path)
    doc.add_paragraph(text)
    doc.save(path)
    print(f"Appended paragraph to {path}")


def merge_docs(paths, out):
    from docx import Document
    from docx.oxml import OxmlElement

    merged = Document(paths[0])
    for path in paths[1:]:
        merged.add_page_break()
        sub = Document(path)
        for element in sub.element.body:
            merged.element.body.append(copy.deepcopy(element))
    merged.save(out)
    print(f"Merged {len(paths)} documents → {out}")
    print("Note: body XML is appended verbatim; styles/numbering are taken from the first document.")


if __name__ == "__main__":
    args = sys.argv[1:]
    if not args or args[0] in {"-h", "--help"}:
        print(__doc__)
        sys.exit(0)

    if "--from-md" in args:
        idx = args.index("--from-md")
        src = args[idx + 1]
        dst = args[idx + 2]
        from_markdown(src, dst)

    elif "--replace" in args:
        idx = args.index("--replace")
        old = args[idx + 1]
        new = args[idx + 2]
        path = args[idx + 3]
        find_replace(old, new, path)

    elif "--append" in args:
        idx = args.index("--append")
        text = args[idx + 1]
        path = args[idx + 2]
        append_paragraph(text, path)

    elif "--merge" in args:
        idx = args.index("--merge")
        # collect paths until --out
        paths = []
        i = idx + 1
        while i < len(args) and args[i] != "--out":
            paths.append(args[i])
            i += 1
        out = args[i + 1] if "--out" in args else "merged.docx"
        merge_docs(paths, out)

    else:
        print("Unknown arguments. Run with --help for usage.")
        sys.exit(1)

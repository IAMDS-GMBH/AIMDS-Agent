#!/usr/bin/env python3
"""Convert Word documents to other formats.

Usage:
    python convert.py document.docx --to pdf
    python convert.py document.docx --to md
    python convert.py document.docx --to html
    python convert.py reports/ --to pdf          # convert entire folder
"""

import subprocess
import sys
import shutil
from pathlib import Path


def _libreoffice_cmd():
    """Return the libreoffice/soffice binary path, or None if not found."""
    for candidate in [
        "libreoffice",
        "soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]:
        if "/" in candidate:
            if Path(candidate).exists():
                return candidate
            continue
        resolved = shutil.which(candidate)
        if resolved:
            return resolved
    return None


def _run_pandoc(src: Path, out: Path, standalone: bool = False) -> tuple[bool, str]:
    cmd = ["pandoc", str(src), "-o", str(out)]
    if standalone:
        cmd.append("--standalone")
    try:
        r = subprocess.run(cmd, capture_output=True, text=True)
    except FileNotFoundError:
        return False, "pandoc not found"
    if r.returncode == 0:
        return True, ""
    return False, (r.stderr or "").strip()


def _markdown_to_basic_html(md_text: str) -> str:
    # Minimal fallback HTML when pandoc is unavailable.
    from html import escape

    body = []
    for line in md_text.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            body.append(f"<h1>{escape(text[2:])}</h1>")
        elif text.startswith("## "):
            body.append(f"<h2>{escape(text[3:])}</h2>")
        elif text.startswith("### "):
            body.append(f"<h3>{escape(text[4:])}</h3>")
        else:
            body.append(f"<p>{escape(text)}</p>")
    return (
        "<!doctype html>\n"
        '<html lang="en">\n<head><meta charset="utf-8"></head>\n<body>\n'
        + "\n".join(body)
        + "\n</body>\n</html>\n"
    )


def _markdown_via_markitdown(src: Path) -> str:
    from markitdown import MarkItDown

    result = MarkItDown().convert(str(src))
    return result.text_content


def _markdown_via_docx(src: Path) -> str:
    from docx import Document

    doc = Document(str(src))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n\n".join(lines)


def _extract_markdown(src: Path) -> str:
    try:
        return _markdown_via_markitdown(src)
    except Exception:
        return _markdown_via_docx(src)


def _to_markdown_with_fallback(src: Path, out: Path) -> None:
    ok, _ = _run_pandoc(src, out)
    if ok:
        print(f"Markdown: {out}")
        return
    out.write_text(_extract_markdown(src), encoding="utf-8")
    print(f"Markdown: {out}")


def _to_html_with_fallback(src: Path, out: Path) -> None:
    ok, _ = _run_pandoc(src, out, standalone=True)
    if ok:
        print(f"HTML: {out}")
        return
    md_text = _extract_markdown(src)
    out.write_text(_markdown_to_basic_html(md_text), encoding="utf-8")
    print(f"HTML: {out}")


def _to_markdown(src: Path, dst: Path | None = None):
    out = dst or src.with_suffix(".md")
    _to_markdown_with_fallback(src, out)


def _to_html(src: Path, dst: Path | None = None):
    out = dst or src.with_suffix(".html")
    _to_html_with_fallback(src, out)


def _to_pdf(src: Path, dst: Path | None = None):
    out = dst or src.with_suffix(".pdf")

    # 1. docx2pdf (macOS/Windows with Word installed)
    try:
        from docx2pdf import convert
        convert(str(src), str(out))
        if out.exists():
            print(f"PDF: {out}")
            return
        print("docx2pdf ran but produced no file (Word may not be installed), trying LibreOffice...", file=sys.stderr)
    except ImportError:
        pass  # not installed
    except Exception as e:
        print(f"docx2pdf failed ({e}), trying LibreOffice...", file=sys.stderr)

    # 2. LibreOffice
    lo = _libreoffice_cmd()
    if lo:
        result = subprocess.run(
            [lo, "--headless", "--convert-to", "pdf", str(src), "--outdir", str(out.parent)],
            capture_output=True, text=True,
        )
        if result.returncode == 0 and out.exists():
            print(f"PDF: {out}")
            return
        print(f"LibreOffice failed: {result.stderr}", file=sys.stderr)

    # 3. fpdf2 text-only fallback (always available if python-docx + fpdf2 installed)
    try:
        _docx_to_pdf_fpdf2(src, out)
        return
    except ImportError as e:
        print(f"fpdf2 fallback unavailable: {e}", file=sys.stderr)

    print(
        "ERROR: No PDF converter found. Install one of:\n"
        "  - docx2pdf  (needs Microsoft Word)  →  uv pip install docx2pdf\n"
        "  - LibreOffice                        →  https://www.libreoffice.org\n"
        "  - fpdf2 + python-docx (text-only)    →  uv pip install fpdf2 python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


def _docx_to_pdf_fpdf2(src: Path, out: Path):
    """Last-resort: extract text from docx and write a basic PDF via fpdf2.
    No images or complex formatting — preserves headings and paragraphs."""
    from docx import Document
    from fpdf import FPDF

    doc = Document(str(src))
    pdf = FPDF()
    pdf.set_margins(left=20, top=20, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(3)
            continue
        sname = para.style.name
        if "Heading 1" in sname:
            pdf.set_font("Helvetica", "B", 16)
            pdf.ln(4)
        elif "Heading 2" in sname:
            pdf.set_font("Helvetica", "B", 13)
            pdf.ln(3)
        elif "Heading 3" in sname:
            pdf.set_font("Helvetica", "B", 11)
            pdf.ln(2)
        else:
            pdf.set_font("Helvetica", size=11)
        # fpdf2 v2+ handles UTF-8 natively — no manual encoding needed
        pdf.multi_cell(0, 6, text)
        pdf.ln(1)

    pdf.output(str(out))
    print(f"PDF (text-only via fpdf2 — install docx2pdf or LibreOffice for full fidelity): {out}")


to_pdf = _to_pdf
to_markdown = _to_markdown
to_html = _to_html


def convert_folder(folder: Path, fmt: str):
    files = list(folder.glob("*.docx"))
    if not files:
        print(f"No .docx files found in {folder}")
        return
    for f in files:
        if fmt == "pdf":
            to_pdf(f)
        elif fmt == "md":
            to_markdown(f)
        elif fmt == "html":
            to_html(f)


if __name__ == "__main__":
    args = sys.argv[1:]
    if not args or args[0] in {"-h", "--help"}:
        print(__doc__)
        sys.exit(0)

    src = Path(args[0]).resolve()
    fmt = "pdf"
    if "--to" in args:
        fmt = args[args.index("--to") + 1].lower()

    if src.is_dir():
        convert_folder(src, fmt)
    elif fmt == "pdf":
        to_pdf(src)
    elif fmt in ("md", "markdown"):
        to_markdown(src)
    elif fmt in ("html", "htm"):
        to_html(src)
    else:
        print(f"Unsupported format: {fmt}. Supported: pdf, md, html")
        sys.exit(1)
